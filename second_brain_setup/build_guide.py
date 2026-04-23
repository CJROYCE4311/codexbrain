"""Generates the 2nd Brain Setup Guide Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, level=1):
    colors = {1: '1F3864', 2: '2E74B5', 3: '2E74B5'}
    sizes  = {1: 20, 2: 15, 3: 13}
    bold   = {1: True, 2: True, 3: True}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = Pt(sizes[level])
    run.font.bold = bold[level]
    r, g, b = bytes.fromhex(colors[level])
    run.font.color.rgb = RGBColor(r, g, b)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_heading(p, level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('NOTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_code(doc, text):
    """Monospace shaded block for terminal commands."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD4, 0x00, 0x00)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_step_box(doc, number, title):
    p = doc.add_paragraph()
    run = p.add_run(f'  STEP {number}:  {title}  ')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F3864')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

# ── Title Page ───────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('2nd Brain Setup Guide')
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('Getting Codex Running on Your Mac\nA Step-by-Step Guide for Non-Technical Users')
sr.font.size = Pt(14)
sr.italic = True
sr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()  # spacer

# ── Overview ─────────────────────────────────────────────────────────────────
add_heading(doc, 'What This Guide Covers', 1)
add_body(doc, 'By the end of this guide you will have everything needed to use your personal 2nd Brain system — a smart, organized way to track your home, finances, health, vehicles, and more using Codex.')

add_body(doc, 'You will install or set up:')
add_bullet(doc, 'A ChatGPT account with Codex access', bold_prefix='ChatGPT account — ')
add_bullet(doc, 'The Codex app — the main interface for managing your 2nd Brain', bold_prefix='Codex app — ')
add_bullet(doc, 'Homebrew — a tool that makes installing Mac software easy', bold_prefix='Homebrew — ')
add_bullet(doc, 'Codex CLI — optional Terminal access when you prefer command-line work', bold_prefix='Codex CLI — ')
add_bullet(doc, 'Visual Studio Code — for browsing your files in a readable format', bold_prefix='Visual Studio Code — ')

add_note(doc, 'This guide is written for older Intel-based MacBooks. The Codex app is available for macOS and Windows; if you are on an Intel Mac, choose the Intel macOS download.')

doc.add_paragraph()

# ── What You Need ─────────────────────────────────────────────────────────────
add_heading(doc, 'Before You Begin — What You Need', 1)
add_bullet(doc, 'Your MacBook (Intel processor is fine)')
add_bullet(doc, 'A working internet connection')
add_bullet(doc, 'An email address for your ChatGPT account')
add_bullet(doc, 'A ChatGPT plan that includes Codex access')
add_bullet(doc, 'About 30–45 minutes of uninterrupted time')

# ── STEP 1 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 1, 'Create or Confirm Your ChatGPT Account')

add_heading(doc, '1a — Sign in at ChatGPT', 2)
add_body(doc, 'Open Safari or Chrome and go to:')
add_code(doc, 'https://chatgpt.com')
add_body(doc, 'Sign in with the account you want to use for Codex. If you do not have an account, create one and follow the prompts.')

add_heading(doc, '1b — Confirm Codex Access', 2)
add_body(doc, 'Codex is included with ChatGPT plans that support Codex. Availability and limits can change, so check the official help page if you are unsure:')
add_code(doc, 'https://help.openai.com/en/articles/11369540')
add_note(doc, 'Use the same ChatGPT account in the Codex app and the Codex CLI so your setup is consistent.')

# ── STEP 2 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 2, 'Open the Terminal')

add_body(doc, 'The Terminal is a text window where you type commands. You will use it for setup, cloning the repository, and optional CLI work.')
add_body(doc, 'To open Terminal:')
add_bullet(doc, 'Press Command (⌘) + Space to open Spotlight Search')
add_bullet(doc, 'Type:  Terminal')
add_bullet(doc, 'Press Enter')
add_body(doc, 'A window with a black or white background and a blinking cursor will appear. This is the Terminal. Leave it open — you will type every command from this guide into it.')
add_note(doc, 'After typing each command, press Enter/Return to run it. Wait for it to finish before typing the next one.')

# ── STEP 3 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 3, 'Install Homebrew')

add_body(doc, 'Homebrew is a free tool that makes it easy to install software on a Mac. Think of it like an App Store for developer tools.')

add_heading(doc, 'Paste this command into Terminal and press Enter:', 2)
add_code(doc, '/bin/bash -c "$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)"')

add_body(doc, 'What to expect:')
add_bullet(doc, 'It will ask for your Mac login password. Type it and press Enter. (You will not see the letters appear — that is normal.)')
add_bullet(doc, 'It may ask you to press Enter again to confirm.')
add_bullet(doc, 'The installation takes 5–15 minutes. Lots of text will scroll by — this is normal.')
add_bullet(doc, 'When it is done, you will see your cursor again with no error messages.')

add_heading(doc, 'Verify Homebrew installed correctly:', 2)
add_code(doc, 'brew --version')
add_body(doc, 'You should see something like:  Homebrew 4.x.x')
add_note(doc, 'If you see "command not found," try closing and reopening Terminal, then try again.')

# ── STEP 4 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 4, 'Install Node.js (Optional, for Codex CLI)')

add_body(doc, 'The Codex app does not require Node.js, but the optional Codex CLI can be installed with npm, which comes with Node.js. Install Node.js with Homebrew:')
add_code(doc, 'brew install node')
add_body(doc, 'When finished, verify:')
add_code(doc, 'node --version')
add_body(doc, 'You should see something like:  v20.x.x or higher.')

# ── STEP 5 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 5, 'Install Codex CLI (Optional Backup)')

add_body(doc, 'Codex CLI is the Terminal version of Codex. Install it by typing:')
add_code(doc, 'npm i -g @openai/codex')
add_body(doc, 'Wait for it to finish (1–3 minutes).')

add_heading(doc, 'Connect Codex CLI to your account:', 2)
add_body(doc, 'Type:')
add_code(doc, 'codex')
add_body(doc, 'The first time you run it, Codex will ask you to sign in. Authenticate with your ChatGPT account and return to Terminal when finished.')
add_body(doc, 'Once you approve it in the browser, switch back to Terminal. You should see a Codex prompt — you are all set.')
add_note(doc, 'To exit Codex CLI at any time, press Ctrl + C or type /exit and press Enter. To update later, run: npm i -g @openai/codex@latest')

# ── STEP 6 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 6, 'Install Visual Studio Code (for Reading Your Files)')

add_body(doc, 'Visual Studio Code (VS Code) lets you open and read your 2nd Brain files in a clean, formatted view. It is free.')
add_body(doc, 'Install it with Homebrew:')
add_code(doc, 'brew install --cask visual-studio-code')
add_body(doc, 'When done, open VS Code:')
add_code(doc, 'open -a "Visual Studio Code"')

add_heading(doc, 'Install the Markdown Preview extension:', 2)
add_body(doc, 'Inside VS Code:')
add_bullet(doc, 'Press Command (⌘) + Shift + X to open Extensions')
add_bullet(doc, 'Search for:  Markdown All in One')
add_bullet(doc, 'Click Install')

add_heading(doc, 'To read a Markdown (.md) file:', 2)
add_bullet(doc, 'Open the file in VS Code (File > Open, or drag the file in)')
add_bullet(doc, 'Press Command (⌘) + Shift + V to see it formatted nicely')

# ── STEP 7 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 7, 'Install the Codex App')

add_body(doc, 'The Codex app is the recommended way to manage this 2nd Brain system. It lets you keep threads, review changes, use worktrees, and sync with GitHub from one desktop app.')
add_body(doc, 'Open the official Codex app docs:')
add_code(doc, 'https://developers.openai.com/codex/app')
add_body(doc, 'Download the app for macOS. If your Mac is Intel-based, choose the Intel macOS build. Open the download and move Codex into Applications if prompted.')
add_body(doc, 'Launch Codex and sign in with the same ChatGPT account from Step 1.')
add_note(doc, 'If macOS blocks the first launch, open System Settings > Privacy & Security and choose Open Anyway for Codex.')

# ── STEP 8 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 8, 'Open Your 2nd Brain Folder in Codex')

add_body(doc, 'Your 2nd Brain files are stored in a folder on your Mac. Codex should be opened on that folder so it can read AGENTS.md and follow the project rules.')

add_heading(doc, 'Get the folder from GitHub:', 2)
add_body(doc, 'In Terminal, type:')
add_code(doc, 'cd ~')
add_code(doc, 'git clone git@github.com:CJROYCE4311/codexbrain.git 2nd_Brain_Codex')
add_body(doc, 'This downloads the folder to your Mac. You only do this once.')

add_heading(doc, 'Open the folder in the Codex app:', 2)
add_bullet(doc, 'Open Codex from Applications')
add_bullet(doc, 'Choose Add project or Open project')
add_bullet(doc, 'Select:  /Users/chrisroyce/2nd_Brain_Codex')
add_bullet(doc, 'Choose Local when you want Codex to edit files on this Mac')
add_body(doc, 'You are now ready. Talk to Codex like you would talk to a person.')

add_heading(doc, 'Optional Terminal start:', 2)
add_code(doc, 'cd ~/2nd_Brain_Codex')
add_code(doc, 'codex')

# ── STEP 9 ────────────────────────────────────────────────────────────────────
add_step_box(doc, 9, 'How to Use Codex for Your 2nd Brain')

add_body(doc, 'Once Codex is open inside your 2nd Brain folder, just type what you want. Here are examples:')

add_heading(doc, 'Adding information:', 2)
add_bullet(doc, '"Add a repair log entry: replaced bathroom faucet on April 20, cost $215, did it myself"')
add_bullet(doc, '"Log an oil change for my 2022 RAV4, today, 38,400 miles, $85 at Jiffy Lube"')
add_bullet(doc, '"Add a new medication: Lisinopril, 10mg, once daily, prescribed by Dr. Jones, started April 1"')
add_bullet(doc, '"Update my grocery budget target to $700/month"')

add_heading(doc, 'Asking questions:', 2)
add_bullet(doc, '"What home repairs have I done this year?"')
add_bullet(doc, '"Show all my current medications"')
add_bullet(doc, '"What\'s my next vehicle maintenance due?"')

add_heading(doc, 'Generating reports:', 2)
add_bullet(doc, '"Create an HTML report of all 2025 expenses"')
add_bullet(doc, '"Summarize my home maintenance history"')

add_heading(doc, 'Ending your session:', 2)
add_body(doc, 'When you are done, tell Codex to save your changes:')
add_bullet(doc, 'Type:  sync changes')
add_bullet(doc, 'Codex will stage, commit, and upload your changes to GitHub')
add_bullet(doc, 'Then close the Codex thread or type /exit if you are using the CLI')

# ── Quick Reference ───────────────────────────────────────────────────────────
add_heading(doc, 'Quick Reference Card', 1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Task'
hdr[1].text = 'What to Type / Do'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    cell._tc.get_or_add_tcPr()

rows = [
    ('Open Terminal',               'Command + Space → type Terminal → Enter'),
    ('Go to your 2nd Brain folder', 'cd ~/2nd_Brain_Codex'),
    ('Start Codex CLI',             'codex'),
    ('Exit Codex CLI',              '/exit or Ctrl + C'),
    ('Open Codex app',              'Applications → Codex → Open project'),
    ('Open a file to read',         'Open VS Code → File > Open → press ⌘+Shift+V'),
    ('Sync/save your changes',      'Tell Codex: "sync changes"'),
    ('Manage your subscription',    'chatgpt.com → Settings'),
]
for task, action in rows:
    row = table.add_row().cells
    row[0].text = task
    row[1].text = action

doc.add_paragraph()

# ── Troubleshooting ───────────────────────────────────────────────────────────
add_heading(doc, 'Troubleshooting', 1)

add_heading(doc, '"command not found" when typing codex', 2)
add_body(doc, 'Close Terminal, reopen it, and try again. If it still fails, run:')
add_code(doc, 'npm i -g @openai/codex')

add_heading(doc, '"Permission denied" errors', 2)
add_body(doc, 'Try adding  sudo  before the command and enter your Mac password when asked.')

add_heading(doc, 'Codex asks to log in again after closing', 2)
add_body(doc, 'This can happen after updates or account changes. Run codex again, or open the Codex app, and follow the sign-in prompt.')

add_heading(doc, 'Browser login window does not open', 2)
add_body(doc, 'Copy the URL shown by Codex in Terminal and paste it into Safari manually, or sign in from the Codex app.')

add_heading(doc, 'Codex app will not open', 2)
add_body(doc, 'Go to System Settings → Privacy & Security, and click "Open Anyway" if macOS blocked it. If the app is too slow or crashes, use Codex CLI in Terminal instead.')

add_heading(doc, 'Git clone fails with SSH permission errors', 2)
add_body(doc, 'Your GitHub SSH key may not be set up on this Mac yet. Ask Codex to help set up GitHub SSH access, or clone with HTTPS if you only need to read the files.')

add_heading(doc, 'You forgot to sync before closing', 2)
add_body(doc, 'Open Terminal, go to the folder, and run:')
add_code(doc, 'cd ~/2nd_Brain_Codex')
add_code(doc, 'git add -A && git commit -m "manual sync" && git push')

# ── Save ──────────────────────────────────────────────────────────────────────
output = '/Users/chrisroyce/2nd_Brain_Codex/second_brain_setup/2nd_Brain_Setup_Guide.docx'
doc.save(output)
print(f'Saved: {output}')
